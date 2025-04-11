from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

def generate_lease_doc(data):
    template_path = "templates/lease_template.docx"
    doc = Document(template_path)
    
    # Replace all placeholders
    for para in doc.paragraphs:
        for key, value in data.items():
            if f'{{{key}}}' in para.text:
                para.text = para.text.replace(f'{{{key}}}', str(value))
    
    # Handle additional clauses
    if 'additional_clauses' in data and data['additional_clauses']:
        # Find the insertion point (before "IN WITNESS")
        insert_index = None
        for i, para in enumerate(doc.paragraphs):
            if "IN WITNESS WHEREOF" in para.text:
                insert_index = i
                break
        
        if insert_index is not None:
            # Get the last clause number (13 in your template)
            last_clause_num = 13
            
            # Insert clauses in reverse order to maintain correct numbering
            for idx, clause in enumerate(reversed(data['additional_clauses']), 1):
                clause_num = last_clause_num + len(data['additional_clauses']) - idx + 1
                
                # Insert the new clause
                new_para = doc.paragraphs[insert_index].insert_paragraph_before(
                    f"{clause_num}. {clause}",
                    style="Normal"
                )
                
                # Format the paragraph
                for run in new_para.runs:
                    run.font.name = 'Cambria'
                    run.font.size = Pt(12)
                
                # Set proper indentation and spacing
                paragraph_format = new_para.paragraph_format
                paragraph_format.line_spacing = 1.5
                paragraph_format.space_after = Pt(12)
                paragraph_format.first_line_indent = Pt(-15)  # Negative indent for numbered list
                paragraph_format.left_indent = Pt(15)         # Positive indent for text
    
    # Set consistent font styling for entire document
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Cambria'
            run.font.size = Pt(12)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(12)
    
    return doc